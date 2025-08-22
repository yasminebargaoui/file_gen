from flask import Flask, request, abort, send_from_directory, jsonify
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Pt
import base64
import os
import uuid

app = Flask(__name__)

# Temporary folder to store generated files
TEMP_DIR = "generated_files"
os.makedirs(TEMP_DIR, exist_ok=True)


# --- Utility functions ---
def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def add_blue_bullet(paragraph, text):
    paragraph.paragraph_format.left_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(2)
    run_bullet = paragraph.add_run("■")
    run_bullet.font.color.rgb = RGBColor(60, 122, 178)
    run_bullet.font.size = Pt(8)
    paragraph.add_run("      ")
    run_text = paragraph.add_run(text)
    run_text.font.size = Pt(11)

def insert_horizontal_line_after(paragraph):
    new_p = insert_paragraph_after(paragraph)
    p = new_p._element
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'dotted')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3399CC')
    pbdr.append(bottom)
    pPr.append(pbdr)
    new_p.paragraph_format.space_before = 1
    new_p.paragraph_format.space_after = Pt(10)
    return new_p


# --- Endpoint to generate DOCX and return link ---
@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    try:
        data = request.get_json()
        if not data:
            return abort(400, "Missing JSON request")

        file_base64 = data.get("file_base64")
        competences = data.get("competences")

        if not file_base64 or not competences:
            return abort(400, "Missing 'file_base64' or 'competences'")

        # Load DOCX from base64
        docx_bytes = base64.b64decode(file_base64)
        doc = Document(BytesIO(docx_bytes))

        # --- Modify document content ---
        # Delete paragraphs between "Connaissances Métier" and "COMPETENCES Projet"
        found_section = False
        paras_to_delete = []
        for para in doc.paragraphs:
            if "Connaissances Métier" in para.text:
                found_section = True
                continue
            if found_section and "COMPETENCES Projet" in para.text:
                break
            if found_section:
                paras_to_delete.append(para)
        for para in paras_to_delete:
            delete_paragraph(para)

        # Find "Connaissances Métier" paragraph
        cm_para = next((p for p in doc.paragraphs if "Connaissances Métier" in p.text), None)
        if not cm_para:
            return abort(400, "'Connaissances Métier' not found")

        cm_para.paragraph_format.space_after = 1
        line_para = insert_horizontal_line_after(cm_para)

        previous_para = line_para
        for comp in competences:
            new_para = insert_paragraph_after(previous_para)
            new_para.paragraph_format.line_spacing = Pt(0)
            add_blue_bullet(new_para, comp)
            previous_para = new_para
        previous_para.paragraph_format.space_after = Pt(9)
        # --- End modification ---

        # Save file with unique name
        file_id = str(uuid.uuid4())
        file_name = f"modified_{file_id}.docx"
        file_path = os.path.join(TEMP_DIR, file_name)
        doc.save(file_path)

        # Return relative download link
        download_url = f"/download/{file_name}"
        return jsonify({"download_url": download_url})

    except Exception as e:
        return abort(500, f"Server error: {str(e)}")


# --- Endpoint to download generated file ---
@app.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(TEMP_DIR, filename, as_attachment=True)


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
